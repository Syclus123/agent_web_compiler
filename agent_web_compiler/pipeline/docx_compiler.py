"""DOCX compilation pipeline.

Uses python-docx to extract content from .docx files.
"""

from __future__ import annotations

import time
from io import BytesIO
from typing import Any

from agent_web_compiler.core.block import Block, BlockType
from agent_web_compiler.core.config import CompileConfig
from agent_web_compiler.core.document import AgentDocument, Quality, SourceType
from agent_web_compiler.core.errors import CompilerError, ParseError

# Heading style name → level mapping.
_HEADING_LEVELS: dict[str, int] = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 4,
    "Heading 5": 5,
    "Heading 6": 6,
}


class DOCXCompiler:
    """Compiles DOCX files into AgentDocuments."""

    def compile(
        self,
        content: bytes,
        *,
        source_file: str | None = None,
        config: CompileConfig | None = None,
    ) -> AgentDocument:
        """Compile DOCX content into an AgentDocument.

        Args:
            content: DOCX content as bytes.
            source_file: Path to the source file.
            config: Compilation configuration.

        Returns:
            An AgentDocument with extracted blocks.

        Raises:
            CompilerError: If python-docx is not installed.
            ParseError: If the DOCX file cannot be parsed.
        """
        try:
            import docx
        except ImportError as exc:
            raise CompilerError(
                "python-docx not installed. Run: pip install 'agent-web-compiler[docx]'",
                stage="docx",
            ) from exc

        if config is None:
            config = CompileConfig()

        timings: dict[str, float] = {}
        t0 = time.perf_counter()

        try:
            doc = docx.Document(BytesIO(content))
        except Exception as e:
            raise ParseError(f"Failed to parse DOCX: {e}", cause=e) from e

        blocks, title, warnings = self._extract_blocks(doc, config)
        timings["extract"] = time.perf_counter() - t0

        # Try to get title from document core properties, fall back to first heading
        if not title:
            try:
                props_title = doc.core_properties.title
                if props_title:
                    title = props_title
            except Exception:
                pass

        # Build canonical markdown
        from agent_web_compiler.exporters.markdown_exporter import to_markdown

        canonical_md = to_markdown(blocks)

        agent_doc = AgentDocument(
            doc_id=AgentDocument.make_doc_id(content),
            source_type=SourceType.DOCX,
            source_file=source_file,
            title=title,
            blocks=blocks,
            canonical_markdown=canonical_md,
            actions=[],
            quality=Quality(
                block_count=len(blocks),
                action_count=0,
                warnings=warnings,
            ),
            debug={"timings": timings} if config.debug else {},
        )

        return agent_doc

    def _extract_blocks(
        self,
        doc: Any,
        config: CompileConfig,
    ) -> tuple[list[Block], str, list[str]]:
        """Walk the DOCX document and extract semantic blocks.

        Args:
            doc: A python-docx Document object.
            config: Compilation configuration.

        Returns:
            Tuple of (blocks, title, warnings).
        """
        warnings: list[str] = []
        blocks: list[Block] = []
        order = 0
        title = ""

        # heading_stack entries: (level, heading_text)
        heading_stack: list[tuple[int, str]] = []

        # --- Walk paragraphs ---
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else "Normal"

            # Determine block type from style
            if style_name in _HEADING_LEVELS:
                level = _HEADING_LEVELS[style_name]
                block_type = BlockType.HEADING

                # Maintain heading stack
                while heading_stack and heading_stack[-1][0] >= level:
                    heading_stack.pop()
                heading_stack.append((level, text))

                if not title:
                    title = text

                section_path = [h[1] for h in heading_stack]
                importance = 0.9

                blocks.append(
                    Block(
                        id=f"b_{order:03d}",
                        type=block_type,
                        text=text,
                        section_path=section_path,
                        order=order,
                        importance=importance,
                        level=level,
                        metadata={"style": style_name},
                    )
                )
                order += 1

            elif "List" in style_name:
                section_path = [h[1] for h in heading_stack]
                blocks.append(
                    Block(
                        id=f"b_{order:03d}",
                        type=BlockType.LIST,
                        text=text,
                        section_path=section_path,
                        order=order,
                        importance=0.6,
                        metadata={"style": style_name},
                    )
                )
                order += 1

            else:
                # Normal paragraph or other style
                section_path = [h[1] for h in heading_stack]
                blocks.append(
                    Block(
                        id=f"b_{order:03d}",
                        type=BlockType.PARAGRAPH,
                        text=text,
                        section_path=section_path,
                        order=order,
                        importance=0.7,
                        metadata={"style": style_name},
                    )
                )
                order += 1

        # --- Walk tables ---
        for table in doc.tables:
            rows_data: list[list[str]] = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                rows_data.append(row_cells)

            if not rows_data:
                continue

            # First row as headers
            headers = rows_data[0] if rows_data else []
            data_rows = rows_data[1:] if len(rows_data) > 1 else []

            text = " | ".join(headers)
            if data_rows:
                for row in data_rows:
                    text += "\n" + " | ".join(row)

            section_path = [h[1] for h in heading_stack]
            blocks.append(
                Block(
                    id=f"b_{order:03d}",
                    type=BlockType.TABLE,
                    text=text,
                    section_path=section_path,
                    order=order,
                    importance=0.8,
                    metadata={
                        "headers": headers,
                        "rows": data_rows,
                        "row_count": len(rows_data),
                        "col_count": len(headers),
                    },
                )
            )
            order += 1

        # --- Detect images (inline shapes) ---
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    blocks.append(
                        Block(
                            id=f"b_{order:03d}",
                            type=BlockType.IMAGE,
                            text="[Image]",
                            section_path=[h[1] for h in heading_stack],
                            order=order,
                            importance=0.4,
                            metadata={"rel_type": rel.reltype},
                        )
                    )
                    order += 1
        except Exception:
            warnings.append("Could not enumerate images from document relationships")

        if not blocks:
            warnings.append("No content blocks extracted from DOCX")

        return blocks, title, warnings
